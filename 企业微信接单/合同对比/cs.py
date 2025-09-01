from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from difflib import Differ


def get_docx_text(file_path):
    """读取 docx 文件内容并返回所有段落的文本"""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return paragraphs


def apply_highlight_to_existing_text(para, diff_line):
    """在原文本上标记差异部分"""
    for run in para.runs:
        if diff_line[2:] in run.text:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # 设置高亮为黄色


def compare_docx_with_highlight(file1, file2, output_file):
    """
    比较两个 docx 文件，并用黄色底色标记差异部分。

    :param file1: 第一个 docx 文件路径
    :param file2: 第二个 docx 文件路径
    :param output_file: 输出结果保存的 docx 文件路径
    """
    # 获取两个文件的文本内容
    text1 = get_docx_text(file1)
    text2 = get_docx_text(file2)

    # 使用 difflib.Differ 进行逐字符对比
    differ = Differ()
    diff = differ.compare(text1, text2)

    # 读取第一个文档并准备修改
    doc = Document(file1)

    para_index = 0
    for line in diff:
        if line.startswith('  '):
            # 没有变化，跳过
            para_index += 1
        elif line.startswith('- '):
            # 在第一个文件中存在，直接在对应段落中标记
            para = doc.paragraphs[para_index]
            apply_highlight_to_existing_text(para, line)
        elif line.startswith('+ '):
            # 在第二个文件中存在的内容，这里忽略
            pass

    # 保存结果
    doc.save(output_file)


if __name__ == "__main__":
    # 输入文件路径
    file1 = "客户修改.docx"
    file2 = "合格模板.docx"
    output_file = "对比文件.docx"

    compare_docx_with_highlight(file1, file2, output_file)
    print(f"带高亮标记的比较结果已保存到 {output_file}")
