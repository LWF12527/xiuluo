from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
import pandas as pd

# 创建文件
d = Document()
# 大标题
d.add_heading("数据分析报告", level=0)


# 标题格式化
def font_song(runs):
    runs.font.name = u'宋体'
    runs._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    runs.font.size = Pt(15)
    runs.font.underline = True


# 段落的全局属性
d.styles['Normal'].font.size = Pt(14)
d.styles['Normal'].font.name = u'仿宋'
d.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 插入表格1
run = d.add_heading("", level=1).add_run("汽车之家型号销售排行榜数据\n")
font_song(run)
table = d.add_table(rows=11, cols=4, style="Table Grid")
table.style.font.color.rgb = RGBColor.from_string("6495ED")  # 字体颜色
data = pd.read_excel(r'data/汽车之家销售排行榜.xlsx', sheet_name="型号销售排行榜")
saleInfo_list = ['排名', '汽车型号', '售价', '型号销售数量']
for i in range(0, 4):  # 属性名
    cell = table.cell(0, i)
    cell.text = saleInfo_list[i]
for i in range(1, 11):  # 行
    for j in range(0, 4):  # 列
        cell = table.cell(i, j)
        cell.text = str(data[saleInfo_list[j]][i - 1])

# 插入表格2
run = d.add_heading("", level=1).add_run("汽车之家品牌销售排行榜数据\n")
font_song(run)
table = d.add_table(rows=11, cols=3, style="Table Grid")
table.style.font.color.rgb = RGBColor.from_string("6495EF")  # 字体颜色
data = pd.read_excel(r'data/汽车之家销售排行榜.xlsx', sheet_name="品牌销售排行榜")
brandInfo_list = ['排名', '品牌名', '品牌销售数量']
for i in range(0, 3):  # 属性名
    cell = table.cell(0, i)
    cell.text = brandInfo_list[i]
for i in range(1, 11):  # 行
    for j in range(0, 3):  # 列
        cell = table.cell(i, j)
        cell.text = str(data[brandInfo_list[j]][i - 1])

run = d.add_heading("", level=1).add_run("品牌销售量占比饼图.png\n")
font_song(run)
d.add_picture("data/品牌销售量占比饼图.png", Inches(7))
d.add_paragraph("   分析得出'比亚迪'品牌的销售量最大，国产汽车开始崛起！")

run = d.add_heading("", level=1).add_run("型号销售量占比饼图.png\n")
font_song(run)
d.add_picture("data/型号销售量占比饼图.png", Inches(7))
d.add_paragraph("   分析得出型号'宋PLUS新能源'销售量最高，'唐新能源'销售量最低，而比亚迪下的销售占比达42%，可见国产新能源汽车非常受欢迎！")

run = d.add_heading("", level=1).add_run("汽车型号销售柱状图.png\n")
font_song(run)
d.add_picture("data/汽车型号销售柱状图.png", Inches(7))

run = d.add_heading("", level=1).add_run("汽车品牌、型号销售柱状图.png\n")
font_song(run)
d.add_picture("data/汽车品牌、型号销售柱状图.png", Inches(7))
# 保存文档
d.save("数据分析报告py.docx")
